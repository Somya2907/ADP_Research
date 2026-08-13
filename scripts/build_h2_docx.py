"""Generate a Word report of Case H2 reasoning traces, phase by phase.

Produces docs/H2_reasoning_traces.docx — a readable report (prose + tables +
full FIRACO traces) rather than slides. Data-driven: traces come from the stored
graphs, L-GED from live embedding alignment, phase numbers from the snapshots and
k-ablation CSVs. Regenerate:

    poetry run python scripts/build_h2_docx.py
"""
from __future__ import annotations

import csv
import os
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src"))
os.environ.setdefault("LEX_DRL_SIMILARITY", "embedding")

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Pt, RGBColor  # noqa: E402

import lex_drl.alignment as alignment_module  # noqa: E402
from lex_drl.alignment import align_all  # noqa: E402
from lex_drl.discrepancy import compute_discrepancies  # noqa: E402
from lex_drl.schema import LegalReasoningGraph  # noqa: E402

GRAPHS = Path("data/outputs/graphs")
OUT = Path("docs/H2_reasoning_traces.docx")
CASE = "H2"

TEACHER_NAVY = RGBColor(0x2F, 0x40, 0x58)
STRONG_GREEN = RGBColor(0x1E, 0x84, 0x49)
WEAK_CLAY = RGBColor(0xB0, 0x69, 0x2B)
FLAG_RED = RGBColor(0x9E, 0x2B, 0x25)


# ── data helpers ──

def _load(name: str) -> LegalReasoningGraph:
    return LegalReasoningGraph.model_validate_json((GRAPHS / name).read_text())


def _norm(s) -> str:
    return " ".join(str(s).split())


def _discrepancy(teacher, student):
    alignment_module.SIMILARITY_METHOD = "embedding"
    return compute_discrepancies(teacher, student, align_all(teacher, student, threshold=0.55))


def _read_csv(path: Path) -> list[dict]:
    return list(csv.DictReader(path.open())) if path.exists() else []


def _snapshot_h2(method: str) -> dict[str, dict]:
    """{student: row} for H2 from a snapshot's discrepancy summary."""
    path = Path(f"data/snapshots/{method}/results/discrepancy_summary.csv")
    out = {}
    for r in _read_csv(path):
        if r["case_id"] == CASE:
            out[r["student"]] = r
    return out


def _rule_citations(name: str) -> list[str]:
    g = _load(name)
    return [f"{r.citation}" for r in g.rules]


# ── docx helpers ──

def h(doc, text, level):
    return doc.add_heading(text, level=level)


def para(doc, text="", *, italic=False, bold=False, size=None, color=None, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, hd in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(hd)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


def trace_section(doc, model, who_label, color):
    """Render one graph's FULL FIRACO trace as headings + bullet lists."""
    p = doc.add_paragraph()
    run = p.add_run(who_label)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = color
    para(doc, f"{model.model_name}  ·  {model.node_count()} nodes  "
              f"(F{len(model.facts)} I{len(model.issues)} R{len(model.rules)} "
              f"A{len(model.applications)} C{len(model.conclusions)} O{len(model.obligations)} "
              f"· E{len(model.edges)})", italic=True, size=9)

    para(doc, f"Facts ({len(model.facts)})", bold=True)
    bullets(doc, [f"{f.fid}: {_norm(f.label)}"
                  + ("" if f.polarity.value == "present" else f"  ({f.polarity.value})")
                  for f in model.facts])

    para(doc, f"Issues ({len(model.issues)})", bold=True)
    bullets(doc, [f"{i.iid}: {_norm(i.label)}  ({i.status.value})" for i in model.issues])

    para(doc, f"Rules ({len(model.rules)})", bold=True)
    bullets(doc, [f"{r.rid}  [{_norm(r.citation)}] — {_norm(r.label)}  "
                  f"({r.authority.value}, {r.jurisdiction})" for r in model.rules])

    para(doc, f"Application ({len(model.applications)} steps)", bold=True)
    bullets(doc, [f"{a.aid} [{a.result.value}] — applies {a.rule_ref} to facts "
                  f"{', '.join(a.fact_refs) if a.fact_refs else '—'} for {a.issue_ref}: "
                  f"{_norm(a.reasoning)}" for a in model.applications])

    para(doc, f"Conclusions ({len(model.conclusions)})", bold=True)
    bullets(doc, [f"{c.cid}: {c.determination.value} / {c.confidence.value}  "
                  f"(supported by {', '.join(c.support_refs) if c.support_refs else '—'})"
                  for c in model.conclusions])

    para(doc, f"Obligations ({len(model.obligations)})", bold=True)
    bullets(doc, [f"{o.oid}: {_norm(o.label)}  ({o.status.value}, {o.jurisdiction}"
                  + (f", by {o.deadline}" if o.deadline else "") + ")"
                  for o in model.obligations])


def lged_annotation(doc, teacher, student):
    rep = _discrepancy(teacher, student)
    para(doc, f"What L-GED counts here — L-GED = {rep.l_ged:.1f}  "
              f"(missed {rep.v_miss_count} · hallucinated {rep.v_halluc_count} · "
              f"edge diffs {rep.e_diff_count} · misgrounded {rep.v_misground_count})",
         bold=True, color=FLAG_RED)
    top_miss = sorted(rep.v_miss, key=lambda m: m.weight, reverse=True)
    if top_miss:
        para(doc, f"Teacher nodes it missed ({len(top_miss)}, highest-cost first):", italic=True)
        bullets(doc, [f"{m.teacher_id} (w{m.weight:g}): {_norm(m.label)}" for m in top_miss[:12]])
        if len(top_miss) > 12:
            para(doc, f"…and {len(top_miss) - 12} more (see docs/REASONING_TRACES.md).", italic=True)
    if rep.v_misground:
        para(doc, f"Right idea, conflicting citation ({len(rep.v_misground)}):", italic=True)
        bullets(doc, [f"{mg.student_id}: cited {_norm(mg.student_citation)} vs teacher "
                      f"{_norm(mg.teacher_citation)} — {_norm(mg.proposition)}"
                      for mg in rep.v_misground])
    return rep


# ── main ──

def main() -> None:
    teacher = _load(f"{CASE}_reference.json")
    gpt5 = _load(f"{CASE}_agent_gpt5.json")
    llama = _load(f"{CASE}_agent_llama3_2b.json")

    tfidf = _snapshot_h2("tfidf_v1")
    embed = _snapshot_h2("embedding_v1")
    dirty = {(r["student"], r["k"]): r for r in _read_csv(Path("results/k_ablation_dirty.csv"))
             if r["case"] == CASE and r["store"] == "dirty"}
    clean = {(r["student"], r["k"]): r for r in _read_csv(Path("results/k_ablation_clean.csv"))
             if r["case"] == CASE and r["store"] == "clean"}

    base_rules = _rule_citations(f"{CASE}_agent_llama3_2b.json")
    dirty_rules = _rule_citations(f"{CASE}_agent_llama3_2b_patched_k3.json")
    clean_rules = _rule_citations(f"{CASE}_agent_llama3_2b_patched_clean_k3.json")

    doc = Document()

    # Title
    title = doc.add_heading("Reasoning Traces — Case H2", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para(doc, "What the teacher and students actually reasoned, and how each project phase "
              "changed it — L-DRL, for the meeting with Prof. Rao.", italic=True)

    # Intro
    h(doc, "What you're looking at", 1)
    para(doc, "Each legal analysis is represented as a graph with six layers — Facts, Issues, "
              "Rules, Application, Conclusion, Obligations (FIRACO) — connected by typed edges. "
              "The teacher (Claude) writes the reference analysis for a case; each student "
              "attempts the same case; the L-GED metric measures how far a student's graph is "
              "from the teacher's (lower = closer = better).")
    para(doc, "Why H2? Of all six cases, H2 shows the largest L-GED movement at every phase — "
              "the biggest metric correction (Phase 1), the biggest apparent patch win "
              "(Phase 2), and the biggest collapse of that win under scrutiny (Phase 3). It is "
              "the clearest single case to see what each phase did.")
    para(doc, "The case: “AutoApprove” is a two-stage AI lending system (Stage 1 fully "
              "automated, Stage 2 human-reviewed). The legal question is whether it is a "
              "high-risk AI system under the Colorado AI Act, whether the vendor is a deployer "
              "or a developer, and what must be fixed by June 30, 2026.")

    # Section 1 — baseline size
    h(doc, "1. The three analyses at a glance", 1)
    para(doc, "Before any intervention, under the trustworthy embedding-based scoring:")
    r_gpt = _discrepancy(teacher, gpt5)
    r_lla = _discrepancy(teacher, llama)
    table(doc,
          ["", "Teacher (Claude)", "GPT-5 (frontier)", "Llama-3B (weak)"],
          [["Total nodes", teacher.node_count(), gpt5.node_count(), llama.node_count()],
           ["Issues", len(teacher.issues), len(gpt5.issues), len(llama.issues)],
           ["Rules", len(teacher.rules), len(gpt5.rules), len(llama.rules)],
           ["Application steps", len(teacher.applications), len(gpt5.applications), len(llama.applications)],
           ["Obligations", len(teacher.obligations), len(gpt5.obligations), len(llama.obligations)],
           ["L-GED (embedding)", "— (reference)", f"{r_gpt.l_ged:.1f}", f"{r_lla.l_ged:.1f}"]])
    bullets(doc, [
        "The teacher builds the full analysis: Stage-1 vs Stage-2 high-risk, deployer-vs-"
        "developer, the memo-vs-statute authority conflict, and a June-2026 compliance-gap "
        "inventory.",
        "GPT-5 tracks that spine (5 of 7 issues) but more compactly.",
        f"Llama collapses the case to {len(llama.issues)} issue and "
        f"{len(llama.applications)} application step — and grounds its lead rule in NYC law "
        "on a Colorado case. Its large L-GED is almost all missing nodes "
        f"({r_lla.v_miss_count}), not wrong ones ({r_lla.v_halluc_count} hallucinations).",
    ])

    # Section 2 — full traces
    h(doc, "2. The full reasoning traces", 1)
    para(doc, "The complete FIRACO analysis each model produced for H2. This is the “show me "
              "the traces” part — every node, full text.")
    h(doc, "2.1 Teacher (Claude) — the reference", 2)
    trace_section(doc, teacher, "Teacher · Claude", TEACHER_NAVY)
    h(doc, "2.2 GPT-5 — frontier student", 2)
    trace_section(doc, gpt5, "GPT-5", STRONG_GREEN)
    lged_annotation(doc, teacher, gpt5)
    h(doc, "2.3 Llama-3B — weak student", 2)
    trace_section(doc, llama, "Llama-3B", WEAK_CLAY)
    lged_annotation(doc, teacher, llama)

    # Section 3 — Phase 1
    h(doc, "3. Phase 1 — Measurement (before → after)", 1)
    para(doc, "Phase 1 built the scoring pipeline and got the measurement right. The key choice "
              "was matching nodes by meaning (sentence embeddings) rather than word overlap. "
              "On H2 this flipped the ranking:")
    table(doc,
          ["H2 L-GED", "GPT-5", "Llama-3B", "Ranking"],
          [["Word overlap (TF-IDF)", tfidf["gpt5"]["l_ged_score"], tfidf["llama3_2b"]["l_ged_score"],
            "GPT-5 scored worse (wrong)"],
           ["Meaning (embeddings)", embed["gpt5"]["l_ged_score"], embed["llama3_2b"]["l_ged_score"],
            "GPT-5 correctly better"]])
    bullets(doc, [
        f"Under TF-IDF, GPT-5's paraphrases of the teacher were counted as hallucinations "
        f"({tfidf['gpt5']['v_halluc_count']} of them), inflating its score above Llama's.",
        f"Embeddings recognize the paraphrases as the same nodes — hallucinations drop to "
        f"{embed['gpt5']['v_halluc_count']}, and the ranking corrects.",
        "This is the project's solid, publishable win: across all 6 cases, embedding ranks "
        "GPT-5 below Llama 6/6; TF-IDF 0/6. H2 is the most dramatic flip.",
    ])

    # Section 4 — Phase 2
    h(doc, "4. Phase 2 — Patch injection (before → after)", 1)
    d = dirty[("llama3_2b", "3")]
    para(doc, "Phase 2 mined reusable “policy patches” from the training cases and injected "
              "the top-k into the students' prompts. On H2 this produced the biggest single win "
              "in the study:")
    table(doc,
          ["Llama H2", "L-GED", "Δ vs baseline"],
          [["No-patch baseline", d["baseline"], "—"],
           ["+ top-3 patches (k=3, dirty store)", d["patched"], f"{float(d['delta']):+.1f}"]])
    para(doc, "What changed in the trace — the rules Llama cited:", bold=True)
    table(doc,
          [f"Baseline ({len(base_rules)} rules)", f"+ patches ({len(dirty_rules)} rules)"],
          [[" · ".join(base_rules), " · ".join(dirty_rules)]])
    bullets(doc, [
        "The patches injected three correct Colorado deployer-duty subsections "
        "(§6-1-1703(2)(a/b/c)) the student had never cited — it looked like genuine rule "
        "recovery.",
        "The hypothesis seemed confirmed: patches help the weak student, best at k=3. So we "
        "pressure-tested it in Phase 3.",
    ])

    # Section 5 — Phase 3
    h(doc, "5. Phase 3 — Data-validity cleanup (before → after)", 1)
    c = clean[("llama3_2b", "3")]
    cg = clean[("gpt5", "3")]
    para(doc, "Phase 3 audited the data, found contamination, quarantined it, pinned one "
              "consistent model run, and re-scored. The H2 win did not survive:")
    table(doc,
          ["Llama H2, k=3", "Phase 2 (dirty)", "Phase 3 (clean)"],
          [["L-GED", d["patched"], c["patched"]],
           ["Δ vs baseline", f"{float(d['delta']):+.1f}", f"{float(c['delta']):+.1f}"],
           ["GPT-5 H2, k=3 (Δ)", f"{float(dirty[('gpt5','3')]['delta']):+.1f}",
            f"{float(cg['delta']):+.1f} (hurt)"]])
    para(doc, "What changed in the trace — the rules Llama cited:", bold=True)
    table(doc,
          [f"Phase 2 dirty ({len(dirty_rules)} rules)", f"Phase 3 clean ({len(clean_rules)} rules)"],
          [[" · ".join(dirty_rules), " · ".join(clean_rules)]])
    bullets(doc, [
        "Two causes of the collapse: (1) the extra Colorado subsections came from patches "
        "mined off a contaminated store and were not independently verified, so quarantine "
        "removed them; (2) the teacher graph itself carries a fabricated §20-875, and GPT-5's "
        "run-to-run variance is ±9.5 — larger than a 4.5 effect.",
        "The clean patch still does one real thing: it fixes Llama's lead-rule grounding from "
        "NYC §20-871 to CO §6-1-1703(2)(a). But it does not add coverage — still one issue.",
        "Honest read: on clean data the patch effect is within noise. Phase 3 was built to "
        "catch exactly this.",
    ])

    # Section 6 — summary
    h(doc, "6. H2 across all three phases", 1)
    table(doc,
          ["Phase", "What it did to H2", "The number"],
          [["Phase 1 — Measurement", "Fixed the ranking (TF-IDF had GPT-5 worse)",
            f"GPT-5 {tfidf['gpt5']['l_ged_score']}→{embed['gpt5']['l_ged_score']}; ranking corrected"],
           ["Phase 2 — Patch injection", "Top-3 patches, Llama k=3",
            f"{float(d['delta']):+.1f} (looked like recovery)"],
           ["Phase 3 — Cleanup", "Quarantine + pinned baseline, Llama k=3",
            f"{float(c['delta']):+.1f} (win mostly contamination)"]])
    bullets(doc, [
        "The metric is trustworthy — H2's ranking corrected, 6/6 overall.",
        "The intervention's headline did not survive honest data hygiene: the recovery was "
        "largely fabricated Colorado subsections plus run-to-run noise.",
        "Next: the ceiling of in-context patching for a 3B model is low — the reasoning trace "
        "still collapses to one issue. The path to real gains is learning from the teacher, "
        "using L-GED as a training reward (see docs/IMPROVEMENT_DIRECTIONS.md).",
    ])

    para(doc, "")
    para(doc, "Full untruncated traces for all cases: docs/REASONING_TRACES.md. Interactive: the "
              "“Reasoning traces” page in scripts/view_graph.py (pick H2). This document "
              "regenerates via scripts/build_h2_docx.py.", italic=True, size=9)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
